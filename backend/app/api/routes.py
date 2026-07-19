"""
API Routes — the AI Firewall's public surface.

Zero-Trust enforcement lives here:
  1. /analyze MUST be called before /send.
  2. /send re-validates that the submitted prompt hash matches (or is a
     masked derivative of) what was analyzed — a client can't skip the
     gate by calling /send directly with an unanalyzed prompt.
  3. LLM responses are re-scanned (Output Guard) before being returned.
"""
import hashlib
import io
import uuid
from typing import Dict

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Header
from sqlalchemy.orm import Session
from pypdf import PdfReader
from docx import Document
import openpyxl

from app.models.schemas import (
    AnalyzeRequest, AnalyzeResponse, SendRequest, SendResponse,
    MaskRequest, MaskResponse, AuditLogEntry, DetectedEntity,
)
from app.analyzers.pipeline import pipeline
from app.analyzers.attachment_scanner import AttachmentScannerAnalyzer
from app.risk.risk_engine import risk_engine
from app.policy.policy_engine import policy_engine
from app.remediation.masker import mask_text
from app.llm.router import llm_router
from app.db.database import get_db
from app.db.models import AuditLog, User
from app.utils.auth import verify_api_key
from app.config import settings

def extract_text_from_file(filename: str, content: bytes) -> str:
    ext = ("." + filename.lower().rsplit(".", 1)[-1]) if "." in filename else ""
    if ext == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(content))
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n".join(text).strip()
        except Exception as e:
            print(f"Error parsing PDF {filename}: {e}")
            return ""
    elif ext == ".docx":
        try:
            doc = Document(io.BytesIO(content))
            text = []
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text.append(" ".join(row_text))
            return "\n".join(text).strip()
        except Exception as e:
            print(f"Error parsing DOCX {filename}: {e}")
            return ""
    elif ext == ".xlsx":
        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            text = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(cell) for cell in row if cell is not None]
                    if row_vals:
                        text.append(" | ".join(row_vals))
            return "\n".join(text).strip()
        except Exception as e:
            print(f"Error parsing XLSX {filename}: {e}")
            return ""
    else:
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return ""

router = APIRouter()

# In-memory session store: request_id -> {prompt_hash, findings, masked_prompt}
# A real deployment would put this in Redis with a TTL.
_ANALYSIS_CACHE: Dict[str, dict] = {}

attachment_scanner = AttachmentScannerAnalyzer()


def _hash_prompt(prompt: str) -> str:
    return hashlib.sha256(prompt.encode()).hexdigest()


def _build_recommendations(findings, violations) -> list[str]:
    recs = []
    for f in findings:
        if not f.passed and f.recommendation:
            recs.append(f.recommendation)
    for v in violations:
        recs.append(f"Policy '{v.policy_name}' triggered: {v.description}")
    # de-dupe, preserve order
    return list(dict.fromkeys(recs))


@router.post("/analyze", response_model=AnalyzeResponse, dependencies=[Depends(verify_api_key)])
async def analyze(
    req: AnalyzeRequest,
    x_user_id: str = Header(default="user1", alias="X-User-Id"),
    db: Session = Depends(get_db)
):
    request_id = str(uuid.uuid4())

    findings = pipeline.run(req.prompt)
    risk = risk_engine.compute(findings)
    policy_result = policy_engine.evaluate(findings)

    decision = policy_result["decision"]
    # Even with no hard policy violation, a very high heuristic risk score blocks.
    if decision == "ALLOW" and risk["score"] >= settings.RISK_BLOCK_THRESHOLD:
        decision = "BLOCK"
    elif decision == "ALLOW" and risk["score"] >= settings.RISK_WARN_THRESHOLD:
        decision = "WARN"

    send_enabled = decision != "BLOCK"

    masked_prompt, masked_count = mask_text(req.prompt, findings)

    _ANALYSIS_CACHE[request_id] = {
        "prompt_hash": _hash_prompt(req.prompt),
        "masked_prompt": masked_prompt,
        "decision": decision,
        "findings": findings,
    }

    entity_types = list(risk["entity_summary"].keys())
    log = AuditLog(
        request_id=request_id,
        decision=decision,
        risk_score=risk["score"],
        risk_level=risk["level"].value,
        entity_types=entity_types,
        policy_violations=[v.policy_name for v in policy_result["violations"]],
        provider=req.provider,
        model=req.model,
        prompt_char_count=len(req.prompt),
        was_sent=False,
        user=x_user_id,
    )
    db.add(log)
    db.commit()

    return AnalyzeResponse(
        request_id=request_id,
        risk_score=risk["score"],
        risk_level=risk["level"],
        decision=decision,
        send_enabled=send_enabled,
        findings=findings,
        policy_violations=policy_result["violations"],
        detected_entity_summary=risk["entity_summary"],
        recommendations=_build_recommendations(findings, policy_result["violations"]),
        masked_prompt_preview=masked_prompt if masked_count > 0 else None,
        auto_fix_available=masked_count > 0,
    )


@router.post("/analyze-file", response_model=AnalyzeResponse, dependencies=[Depends(verify_api_key)])
async def analyze_file(
    prompt: str = Form(""),
    file: UploadFile = File(...),
    x_user_id: str = Header(default="user1", alias="X-User-Id"),
    db: Session = Depends(get_db)
):
    content = await file.read()
    if len(content) > settings.MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(413, f"File exceeds {settings.MAX_UPLOAD_MB}MB limit.")

    file_finding = attachment_scanner.scan_file(file.filename, content)

    extracted_text = extract_text_from_file(file.filename, content)

    combined_text = f"{prompt}\n{extracted_text}" if prompt else extracted_text
    findings = pipeline.run(combined_text, filename=file.filename)
    findings.append(file_finding)

    risk = risk_engine.compute(findings)
    policy_result = policy_engine.evaluate(findings)
    decision = policy_result["decision"]
    if decision == "ALLOW" and risk["score"] >= settings.RISK_BLOCK_THRESHOLD:
        decision = "BLOCK"
    elif decision == "ALLOW" and risk["score"] >= settings.RISK_WARN_THRESHOLD:
        decision = "WARN"

    request_id = str(uuid.uuid4())
    masked_prompt, masked_count = mask_text(combined_text, findings)
    _ANALYSIS_CACHE[request_id] = {
        "prompt_hash": _hash_prompt(combined_text),
        "masked_prompt": masked_prompt,
        "decision": decision,
        "findings": findings,
    }

    entity_types = list(risk["entity_summary"].keys())
    db.add(AuditLog(
        request_id=request_id, decision=decision, risk_score=risk["score"],
        risk_level=risk["level"].value, entity_types=entity_types,
        policy_violations=[v.policy_name for v in policy_result["violations"]],
        prompt_char_count=len(combined_text), was_sent=False,
        user=x_user_id,
    ))
    db.commit()

    return AnalyzeResponse(
        request_id=request_id,
        risk_score=risk["score"],
        risk_level=risk["level"],
        decision=decision,
        send_enabled=decision != "BLOCK",
        findings=findings,
        policy_violations=policy_result["violations"],
        detected_entity_summary=risk["entity_summary"],
        recommendations=_build_recommendations(findings, policy_result["violations"]),
        masked_prompt_preview=masked_prompt if masked_count > 0 else None,
        auto_fix_available=masked_count > 0,
        extracted_text=extracted_text,
        document_classifications=[{"filename": file.filename}],
    )


@router.post("/mask", response_model=MaskResponse, dependencies=[Depends(verify_api_key)])
async def mask(req: MaskRequest):
    findings = pipeline.run(req.prompt)
    masked, count = mask_text(req.prompt, findings)
    return MaskResponse(masked_prompt=masked, entities_masked=count)


@router.post("/send", response_model=SendResponse, dependencies=[Depends(verify_api_key)])
async def send(req: SendRequest, db: Session = Depends(get_db)):
    cached = _ANALYSIS_CACHE.get(req.request_id)
    if not cached:
        raise HTTPException(400, "No prior /analyze call found for this request_id. Zero-Trust policy requires analysis before send.")

    if cached["decision"] == "BLOCK" and not req.use_masked_version:
        raise HTTPException(403, "This prompt was BLOCKED by policy. Apply masking or revise the prompt before sending.")

    outbound_prompt = cached["masked_prompt"] if req.use_masked_version else req.prompt

    # Re-verify: if using the raw (non-masked) prompt, it must match what was analyzed.
    if not req.use_masked_version and _hash_prompt(req.prompt) != cached["prompt_hash"]:
        raise HTTPException(409, "Prompt has changed since analysis. Please re-run Analyze before sending.")

    try:
        response_text, provider_used, model_used = await llm_router.route(
            outbound_prompt, req.provider, req.model
        )
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(502, str(exc))

    # Output Guard: scan the LLM's response before returning it to the user.
    output_findings = pipeline.run(response_text)
    flagged = [f for f in output_findings if not f.passed]
    was_sanitized = False
    if flagged:
        response_text, _ = mask_text(response_text, output_findings)
        was_sanitized = True

    log = db.query(AuditLog).filter(AuditLog.request_id == req.request_id).first()
    if log:
        log.was_sent = True
        log.provider = provider_used
        log.model = model_used
        db.commit()

    return SendResponse(
        response_text=response_text,
        provider_used=provider_used,
        model_used=model_used,
        output_flags=flagged,
        was_sanitized=was_sanitized,
    )


@router.get("/audit-logs", dependencies=[Depends(verify_api_key)])
async def get_audit_logs(
    limit: int = 50,
    x_user_id: str = Header(default="user1", alias="X-User-Id"),
    db: Session = Depends(get_db)
):
    query = db.query(AuditLog)
    user_record = db.query(User).filter(User.email == x_user_id.strip().lower()).first()
    is_admin = (x_user_id.lower() == "admin") or (user_record and user_record.role == "security_admin")
    if not is_admin:
        query = query.filter(AuditLog.user == x_user_id)
    logs = query.order_by(AuditLog.timestamp.desc()).limit(limit).all()
    return [
        {
            "id": l.id, "request_id": l.request_id, "timestamp": l.timestamp,
            "decision": l.decision, "risk_score": l.risk_score, "risk_level": l.risk_level,
            "entity_types": l.entity_types, "policy_violations": l.policy_violations,
            "provider": l.provider, "model": l.model, "was_sent": l.was_sent,
        }
        for l in logs
    ]


@router.get("/dashboard/summary", dependencies=[Depends(verify_api_key)])
async def dashboard_summary(
    x_user_id: str = Header(default="user1", alias="X-User-Id"),
    db: Session = Depends(get_db)
):
    query = db.query(AuditLog)
    user_record = db.query(User).filter(User.email == x_user_id.strip().lower()).first()
    is_admin = (x_user_id.lower() == "admin") or (user_record and user_record.role == "security_admin")
    if not is_admin:
        query = query.filter(AuditLog.user == x_user_id)
    logs = query.all()
    total = len(logs)
    blocked = sum(1 for l in logs if l.decision == "BLOCK")
    warned = sum(1 for l in logs if l.decision == "WARN")
    allowed = sum(1 for l in logs if l.decision == "ALLOW")

    entity_counts: Dict[str, int] = {}
    policy_counts: Dict[str, int] = {}
    for l in logs:
        for et in (l.entity_types or []):
            entity_counts[et] = entity_counts.get(et, 0) + 1
        for pv in (l.policy_violations or []):
            policy_counts[pv] = policy_counts.get(pv, 0) + 1

    return {
        "requests_processed": total,
        "blocked": blocked,
        "warned": warned,
        "allowed": allowed,
        "top_detected_entities": sorted(entity_counts.items(), key=lambda x: -x[1])[:10],
        "top_policy_violations": sorted(policy_counts.items(), key=lambda x: -x[1])[:10],
    }


@router.get("/policies", dependencies=[Depends(verify_api_key)])
async def get_policies():
    return policy_engine.list_policies()


@router.put("/policies", dependencies=[Depends(verify_api_key)])
async def update_policies(policies: list[dict]):
    policy_engine.save_policies(policies)
    return {"status": "saved", "count": len(policies)}


@router.get("/analysis/report", dependencies=[Depends(verify_api_key)])
async def get_analysis_report(
    start_date: str,
    end_date: str,
    x_user_id: str = Header(default="user1", alias="X-User-Id"),
    db: Session = Depends(get_db)
):
    from datetime import date, datetime, time
    try:
        start_dt = datetime.combine(date.fromisoformat(start_date), time.min)
        end_dt = datetime.combine(date.fromisoformat(end_date), time.max)
    except ValueError as e:
        raise HTTPException(400, f"Invalid date format. Expected YYYY-MM-DD. Error: {e}")

    query = db.query(AuditLog).filter(
        AuditLog.timestamp >= start_dt,
        AuditLog.timestamp <= end_dt
    )
    user_record = db.query(User).filter(User.email == x_user_id.strip().lower()).first()
    is_admin = (x_user_id.lower() == "admin") or (user_record and user_record.role == "security_admin")
    if not is_admin:
        query = query.filter(AuditLog.user == x_user_id)
    logs = query.all()

    total = len(logs)
    if total == 0:
        return {
            "total_requests": 0,
            "decisions": {"ALLOW": 0, "WARN": 0, "BLOCK": 0},
            "avg_risk_score": 0.0,
            "avg_prompt_char_count": 0.0,
            "was_sent_percentage": 0.0,
            "daily_trends": [],
            "entity_counts": {},
            "policy_counts": {},
            "provider_counts": {},
            "risk_level_counts": {"LOW": 0, "MEDIUM": 0, "HIGH": 0, "CRITICAL": 0},
            "insights": ["No audit logs found for the selected date range."]
        }

    blocked = sum(1 for l in logs if l.decision == "BLOCK")
    warned = sum(1 for l in logs if l.decision == "WARN")
    allowed = sum(1 for l in logs if l.decision == "ALLOW")
    
    avg_risk = sum(l.risk_score for l in logs) / total
    avg_char = sum(l.prompt_char_count for l in logs) / total
    was_sent_count = sum(1 for l in logs if l.was_sent)
    was_sent_pct = (was_sent_count / total) * 100

    entity_counts = {}
    policy_counts = {}
    provider_counts = {}
    risk_level_counts = {"LOW": 0, "MEDIUM": 0, "HIGH": 0, "CRITICAL": 0}
    
    for l in logs:
        for et in (l.entity_types or []):
            entity_counts[et] = entity_counts.get(et, 0) + 1
        for pv in (l.policy_violations or []):
            policy_counts[pv] = policy_counts.get(pv, 0) + 1
        
        prov_model = f"{l.provider or 'Unknown'}/{l.model or 'Unknown'}"
        provider_counts[prov_model] = provider_counts.get(prov_model, 0) + 1
        
        rl = l.risk_level or "LOW"
        risk_level_counts[rl] = risk_level_counts.get(rl, 0) + 1

    daily_data = {}
    for l in logs:
        day_str = l.timestamp.date().isoformat()
        if day_str not in daily_data:
            daily_data[day_str] = {"count": 0, "total_risk": 0}
        daily_data[day_str]["count"] += 1
        daily_data[day_str]["total_risk"] += l.risk_score

    daily_trends = []
    for d_str in sorted(daily_data.keys()):
        daily_trends.append({
            "date": d_str,
            "count": daily_data[d_str]["count"],
            "avg_risk": round(daily_data[d_str]["total_risk"] / daily_data[d_str]["count"], 1)
        })

    insights = []
    if daily_trends:
        peak_day = max(daily_trends, key=lambda x: x["count"])
        insights.append(f"Traffic peaked on {peak_day['date']} with {peak_day['count']} requests analyzed.")

    blocked_logs = [l for l in logs if l.decision == "BLOCK"]
    allowed_logs = [l for l in logs if l.decision == "ALLOW"]
    if blocked_logs and allowed_logs:
        avg_blocked_len = sum(l.prompt_char_count for l in blocked_logs) / len(blocked_logs)
        avg_allowed_len = sum(l.prompt_char_count for l in allowed_logs) / len(allowed_logs)
        diff_pct = ((avg_blocked_len - avg_allowed_len) / (avg_allowed_len or 1)) * 100
        if diff_pct > 0:
            insights.append(f"Blocked prompts are on average {diff_pct:.1f}% longer ({int(avg_blocked_len)} chars) than allowed prompts ({int(avg_allowed_len)} chars).")
        else:
            insights.append(f"Allowed prompts are on average {abs(diff_pct):.1f}% longer ({int(avg_allowed_len)} chars) than blocked prompts ({int(avg_blocked_len)} chars).")

    if provider_counts:
        top_model = max(provider_counts.items(), key=lambda x: x[1])[0]
        insights.append(f"The most active target LLM provider/model is {top_model} ({provider_counts[top_model]} requests).")

    if entity_counts:
        top_entity, top_ent_count = max(entity_counts.items(), key=lambda x: x[1])
        insights.append(f"Sensitive entity '{top_entity}' was the most frequent trigger, flagged in {top_ent_count} instances.")

    total_violations = sum(policy_counts.values())
    if total_violations > 0:
        insights.append(f"A total of {total_violations} policy rules were violated in this date range.")

    return {
        "total_requests": total,
        "decisions": {
            "ALLOW": allowed,
            "WARN": warned,
            "BLOCK": blocked
        },
        "avg_risk_score": round(avg_risk, 1),
        "avg_prompt_char_count": round(avg_char, 1),
        "was_sent_percentage": round(was_sent_pct, 1),
        "daily_trends": daily_trends,
        "entity_counts": entity_counts,
        "policy_counts": policy_counts,
        "provider_counts": provider_counts,
        "risk_level_counts": risk_level_counts,
        "insights": insights
    }


@router.get("/health")
async def health():
    return {"status": "ok", "service": "SecurePrompt AI Gateway"}


# --- Authentication & User Management ---
import os
import hashlib
from pydantic import BaseModel

def hash_password(password: str) -> str:
    salt = os.urandom(16)
    hashed = hashlib.scrypt(
        password.encode('utf-8'),
        salt=salt,
        n=16384,
        r=8,
        p=1
    )
    return f"{salt.hex()}:{hashed.hex()}"

def verify_password(password: str, stored_hash: str) -> bool:
    try:
        salt_hex, hash_hex = stored_hash.split(":")
        salt = bytes.fromhex(salt_hex)
        computed = hashlib.scrypt(
            password.encode('utf-8'),
            salt=salt,
            n=16384,
            r=8,
            p=1
        )
        return computed.hex() == hash_hex
    except Exception:
        return False

class SignUpRequest(BaseModel):
    name: str
    email: str
    password: str
    role: str  # 'developer' or 'security_admin'

class LoginRequest(BaseModel):
    email: str
    password: str

@router.post("/auth/signup")
async def signup(req: SignUpRequest, db: Session = Depends(get_db)):
    # Check if user already exists
    existing = db.query(User).filter(User.email == req.email.strip().lower()).first()
    if existing:
        raise HTTPException(400, "A user with this email already exists.")

    hashed = hash_password(req.password)
    user = User(
        name=req.name.strip(),
        email=req.email.strip().lower(),
        hashed_password=hashed,
        role=req.role
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {
        "id": user.id,
        "name": user.name,
        "email": user.email,
        "role": user.role
    }

@router.post("/auth/login")
async def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == req.email.strip().lower()).first()
    if not user:
        raise HTTPException(400, "Invalid email or password.")

    if not verify_password(req.password, user.hashed_password):
        raise HTTPException(400, "Invalid email or password.")

    return {
        "id": user.id,
        "name": user.name,
        "email": user.email,
        "role": user.role
    }
